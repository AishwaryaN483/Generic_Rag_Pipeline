from docx import Document as DocxDocument
from langchain_core.documents import Document


def parse_docx(file_path):

    doc = DocxDocument(file_path)

    text = ""

    for para in doc.paragraphs:
        text += para.text + "\n"

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path,
                "type": "docx"
            }
        )
    ]