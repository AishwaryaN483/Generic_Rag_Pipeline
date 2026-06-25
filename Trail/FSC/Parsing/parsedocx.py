# from docx import Document

# def parse_docx(file_path):
#     doc = Document(file_path)
#     text_parts = []
#     table_texts = []

#     for para in doc.paragraphs:
#         text_parts.append(para.text)

#     for table in doc.tables:
#         rows = []
#         for row in table.rows:
#             rows.append(", ".join(cell.text.strip() for cell in row.cells))
#         table_texts.append("\n".join(rows))

#     combined_text = "\n".join(text_parts)
#     return {"text": combined_text, "tables": table_texts}


from docx import Document
import re

def clean_text(text):
    text = text.strip()

    # Normalize spaces
    text = re.sub(r"\s+", " ", text)

    # Fix OCR repeated characters (EEEE → E)
    text = re.sub(r'(.)\1{3,}', r'\1', text)

    # Remove non-ascii garbage (but keep common symbols)
    text = re.sub(r'[^\x00-\x7F]+', '', text)

    # Remove placeholder junk
    if re.search(r'_+', text):
        return ""

    # Remove lines with almost no letters (pure noise)
    alpha_ratio = sum(c.isalpha() for c in text) / (len(text) + 1)

    if alpha_ratio < 0.2:
        return ""

    return text


def remove_duplicates_preserve_order(text_list):
    seen = set()
    result = []
    for item in text_list:
        if item not in seen and item != "":
            seen.add(item)
            result.append(item)
    return result


def parse_docx(file_path):
    doc = Document(file_path)

    full_text_parts = []
    structured_data = []
    global_seen = set()

    for para in doc.paragraphs:
        text = clean_text(para.text)

        if not text or text in global_seen:
            continue

        global_seen.add(text)

        # Detect heading
        style = para.style.name.lower()
        if "heading" in style:
            tag = "heading"
            formatted = f"[HEADING] {text}"
        else:
            tag = "paragraph"
            formatted = text

        structured_data.append({
            "type": tag,
            "text": text
        })

        full_text_parts.append(formatted)

    # -------- TABLE EXTRACTION --------
    table_data = []

    for table in doc.tables:
        rows = []

        for row in table.rows:
            row_text = " | ".join(
                clean_text(cell.text) for cell in row.cells if clean_text(cell.text)
            )

            if row_text:
                rows.append(row_text)

        if rows:
            table_text = "\n".join(rows)

            if table_text not in global_seen:
                table_data.append(table_text)
                global_seen.add(table_text)
                full_text_parts.append(table_text)

    combined_text = "\n\n".join(full_text_parts).strip()

    return {
        "doc_id": file_path,
        "text": combined_text,
        "structure": structured_data,
        "tables": table_data,
        "metadata": {
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "source": "docx"
        }
    }


# -------- RUN --------
if __name__ == "__main__":
    file_path = "D://topicTrace//Trail//FSC//Samples//docx//Sample2.docx"

    result = parse_docx(file_path)

    print("Paragraphs:", result["metadata"]["total_paragraphs"])
    print("Tables:", result["metadata"]["total_tables"])
    print("\nPreview:\n")
    print(result["text"][:500])

    with open("parsed_docx_output.txt", "w", encoding="utf-8") as f:
        f.write(result["text"])

    print("\nSaved to parsed_docx_output.txt")